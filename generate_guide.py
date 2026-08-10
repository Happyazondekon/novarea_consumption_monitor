from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_guide():
    doc = Document()

    # Style: Title
    title = doc.add_heading('GUIDE D\'UTILISATION : PLATEFORME DE MONITORING NOVAREA TEXTILES BENIN 🏭📊', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info block
    p = doc.add_paragraph()
    p.add_run('Version : 1.0\n').bold = True
    p.add_run('Statut : Document Industriel Final\n').bold = True
    p.add_run('Date : 01/08/2026').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph('_' * 50)

    # Section 1
    doc.add_heading('1. INTRODUCTION ET ACCÈS', level=1)
    doc.add_paragraph('Ce guide détaille l\'utilisation de la plateforme intégrée de suivi des consommations (Électricité et Eau) pour l\'usine Novarea Textiles. La plateforme est optimisée pour une utilisation sur Desktop (Audit) et Mobile (Saisie Terrain).')

    doc.add_heading('1.1 Accès à la Plateforme', level=2)
    p = doc.add_paragraph('L\'application est accessible via l\'URL sécurisée : ')
    p.add_run('https://novarea-consumption.vercel.app').bold = True

    # Placeholder for Capture
    p = doc.add_paragraph('\n[CADRE POUR CAPTURE : ÉCRAN DE LOGIN DESKTOP AVEC BRANDING]\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1.2 Installation Mobile (PWA)', level=2)
    doc.add_paragraph('Pour installer l\'application sur votre smartphone comme une application native :')
    doc.add_paragraph('1. Ouvrez l\'URL dans Safari (iOS) ou Chrome (Android).', style='List Number')
    doc.add_paragraph('2. Cliquez sur "Partager" ou le menu "Options".', style='List Number')
    doc.add_paragraph('3. Sélectionnez "Sur l\'écran d\'accueil" ou "Installer l\'application".', style='List Number')

    p = doc.add_paragraph('\n[CADRE POUR CAPTURE : INSTALLATION PWA SUR MOBILE]\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('_' * 50)

    # Section 2
    doc.add_heading('2. AUTHENTIFICATION SÉCURISÉE', level=1)
    doc.add_paragraph('Le système utilise une gestion des accès basée sur les rôles (RBAC).')
    doc.add_paragraph('1. Saisissez votre Identifiant.', style='List Number')
    doc.add_paragraph('2. Saisissez votre Mot de passe.', style='List Number')
    doc.add_paragraph('3. Cliquez sur LOGIN.', style='List Number')
    doc.add_paragraph('Note : La session est limitée à 1 heure pour des raisons de sécurité industrielle.', style='Intense Quote')

    doc.add_paragraph('_' * 50)

    # Section 3
    doc.add_heading('3. PROFIL : ADMINISTRATEUR (HUB ANALYTIQUE)', level=1)
    doc.add_paragraph('L\'administrateur dispose d\'une vue d\'ensemble sur toute l\'infrastructure.')

    doc.add_heading('3.1 Tableau de Bord (Dashboard)', level=2)
    doc.add_paragraph('Le dashboard présente les KPIs en temps réel :')
    doc.add_paragraph('Consommations du jour (Électricité & Eau).', style='List Bullet')
    doc.add_paragraph('Moyennes journalières calculées sur les 30 derniers jours.', style='List Bullet')
    doc.add_paragraph('Graphiques Dynamiques : Analyse croisée entre la consommation et les événements opérationnels.', style='List Bullet')

    p = doc.add_paragraph('\n[CADRE POUR CAPTURE : DASHBOARD ADMINISTRATEUR - VUE ANALYTIQUE]\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.2 Audit des Soumissions (Submission Audit)', level=2)
    doc.add_paragraph('Toutes les saisies faites par les techniciens doivent être auditées ici.')
    doc.add_paragraph('Validation : Validez les index capturés.', style='List Bullet')
    doc.add_paragraph('Édition : Corrigez une erreur de saisie si nécessaire.', style='List Bullet')
    doc.add_paragraph('Preuve Visuelle : Visualisez la photo du compteur jointe à chaque lecture.', style='List Bullet')
    doc.add_paragraph('Actions Groupées : Possibilité de valider ou supprimer plusieurs lignes simultanément.', style='List Bullet')

    p = doc.add_paragraph('\n[CADRE POUR CAPTURE : LISTE D\'AUDIT AVEC MODALE DE PRÉVISUALISATION IMAGE]\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.3 Contexte Journalier (Daily Events)', level=2)
    doc.add_paragraph('C\'est ici que l\'administrateur enregistre les anomalies de production (Coupures, Fuites, Pics PTR, etc.) pour expliquer les variations de consommation sur les graphiques.')

    p = doc.add_paragraph('\n[CADRE POUR CAPTURE : WIZARD DE CRÉATION D\'ÉVÉNEMENT EN 3 ÉTAPES]\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.4 Centre de Génération de Rapports', level=2)
    doc.add_paragraph('Outil de construction de documents professionnels (PDF & Excel).')
    doc.add_paragraph('Filtres : Semaine, Mois, Année.', style='List Bullet')
    doc.add_paragraph('Export PDF : Rapport haute-fidélité avec graphiques capturés et tableaux de données.', style='List Bullet')
    doc.add_paragraph('Export Excel : Export multi-feuilles (Données et Événements séparés).', style='List Bullet')

    p = doc.add_paragraph('\n[CADRE POUR CAPTURE : GÉNÉRATEUR DE RAPPORT AVEC APERÇU LIVE]\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('_' * 50)

    # Section 4
    doc.add_heading('4. PROFIL : ÉLECTRICIEN / TECHNICIEN (HUB OPÉRATIONNEL)', level=1)
    doc.add_paragraph('Le technicien utilise principalement l\'interface mobile pour ses relevés.')

    doc.add_heading('4.1 Accueil Mobile', level=2)
    doc.add_paragraph('Accès rapide aux indicateurs de performance de son shift et aux dernières instructions.')

    p = doc.add_paragraph('\n[CADRE POUR CAPTURE : DASHBOARD TECHNICIEN SUR MOBILE]\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('4.2 Enregistrement d\'un Relevé (New Reading)', level=2)
    doc.add_paragraph('Processus simplifié en 3 étapes :')
    doc.add_paragraph('1. Ressource : Choisir Électricité ou Eau.', style='List Number')
    doc.add_paragraph('2. Index : Saisir la valeur du compteur (Boutons +/- pour ajustement rapide).', style='List Number')
    doc.add_paragraph('3. Photo : Capturer obligatoirement une photo du compteur comme preuve.', style='List Number')

    p = doc.add_paragraph('\n[CADRE POUR CAPTURE : WORKFLOW NEW READING SUR MOBILE]\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('4.3 Gestion des Missions (Mission Control)', level=2)
    doc.add_paragraph('Le technicien voit les directives envoyées par l\'administrateur. Une fois la tâche terminée, il clique sur "DONE". L\'administrateur verra instantanément le statut mis à jour.')

    p = doc.add_paragraph('\n[CADRE POUR CAPTURE : LISTE DES MISSIONS AVEC STATUT PENDING/DONE]\n')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('_' * 50)

    # Section 5
    doc.add_heading('5. PARAMÈTRES ET SÉCURITÉ', level=1)
    doc.add_paragraph('Accessible via l\'icône de profil en haut à droite.')
    doc.add_paragraph('Mon Profil : Mise à jour du nom, de l\'email et de la photo de profil.', style='List Bullet')
    doc.add_paragraph('Sécurité : Changement du mot de passe.', style='List Bullet')
    doc.add_paragraph('Gestion d\'Équipe (Admin uniquement) : Création et révocation des accès pour le personnel.', style='List Bullet')

    doc.add_paragraph('\n' * 2)
    p = doc.add_paragraph('Document édité pour Novarea Textiles Benin. 🛡️🏭')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save('GUIDE_UTILISATION_NOVAREA.docx')
    print("DOCX generated successfully.")

if __name__ == "__main__":
    create_guide()
